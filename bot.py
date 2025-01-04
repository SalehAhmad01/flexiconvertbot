import os
import tempfile
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    ApplicationBuilder, CommandHandler, MessageHandler, filters, CallbackQueryHandler, ContextTypes
)
from PIL import Image
from gtts import gTTS
import pytesseract
from fpdf import FPDF
from docx import Document
from pdf2docx import Converter
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Tesseract configuration
pytesseract.pytesseract.tesseract_cmd = os.getenv('TESSERACT_PATH', '/usr/bin/tesseract')

# Telegram Bot Token
TOKEN = os.getenv('BOT_TOKEN')
if not TOKEN:
    raise ValueError("BOT_TOKEN is missing in environment variables.")

# Set up Application
app = ApplicationBuilder().token(TOKEN).build()

# Conversion functions
def docx_to_pdf(input_path, output_path):
    """Convert DOCX to PDF using the FPDF library."""
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Input file not found: {input_path}")
    
    try:
        doc = Document(input_path)
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                pdf.multi_cell(0, 10, paragraph.text)

        pdf.output(output_path)
    except Exception as e:
        print(f"Error converting DOCX to PDF: {e}")
        raise

def pdf_to_docx(input_path, output_path):
    """Convert PDF to DOCX using pdf2docx."""
    try:
        converter = Converter(input_path)
        converter.convert(output_path, start=0, end=None)
        converter.close()
    except Exception as e:
        print(f"Error converting PDF to DOCX: {e}")
        raise

def images_to_pdf(image_paths, output_path):
    """Convert images to PDF."""
    try:
        images = [Image.open(img).convert('RGB') for img in image_paths if img.lower().endswith(('jpg', 'jpeg', 'png'))]
        if images:
            images[0].save(output_path, save_all=True, append_images=images[1:])
        else:
            print("No valid images found for PDF conversion.")
    except Exception as e:
        print(f"Error converting images to PDF: {e}")
        raise

def image_to_text(image_path):
    """Extract text from an image."""
    try:
        with Image.open(image_path) as img:
            img = img.convert('L')
            text = pytesseract.image_to_string(img)
            return text.strip()
    except Exception as e:
        print(f"Error extracting text from image: {e}")
        return "Failed to extract text from the image."

def text_to_speech(text, output_path):
    """Convert text to speech."""
    try:
        tts = gTTS(text=text, lang='en')
        tts.save(output_path)
    except Exception as e:
        print(f"Error converting text to speech: {e}")
        raise

def text_to_pdf(text, output_path):
    """Convert text to PDF."""
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for line in text.split('\n'):
            pdf.multi_cell(0, 10, txt=line, align='L')
        pdf.output(output_path)
    except Exception as e:
        print(f"Error converting text to PDF: {e}")
        raise

def cleanup(file_path):
    """Delete temporary files."""
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
        except Exception as e:
            print(f"Error cleaning up file {file_path}: {e}")


# Telegram Bot Handlers
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle /start command."""
    greeting_text = (
        "Hello! I am a bot that can:\n"
        "- Convert documents to PDF\n"
        "- Extract text from images\n"
        "- Convert text to speech\n"
        "- Convert PDF to text\n"
        "Upload a file or send text to start!"
    )
    await update.message.reply_text(greeting_text)

async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle text messages."""
    text = update.message.text
    context.user_data['text'] = text
    print(f"Received text: {text}")  # Debugging line
    keyboard = [
        [InlineKeyboardButton("Speech", callback_data="Speech")],
        [InlineKeyboardButton("PDF", callback_data="PDF")],
        [InlineKeyboardButton("DOCX", callback_data="DOCX")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text("Please choose a conversion option:", reply_markup=reply_markup)

async def handle_image(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle image messages."""
    photo = update.message.photo[-1]
    file = await photo.get_file()
    file_path = os.path.join(tempfile.gettempdir(), "image.jpg")
    await file.download_to_drive(file_path)

    text = image_to_text(file_path)
    if text:
        await update.message.reply_text(f"Extracted text: {text}")
        context.user_data['text'] = text
    else:
        await update.message.reply_text("No text detected in the image.")
    cleanup(file_path)

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle document uploads."""
    document = update.message.document
    file = await document.get_file()
    file_path = os.path.join(tempfile.gettempdir(), document.file_name)
    await file.download_to_drive(file_path)

    context.user_data['file_path'] = file_path

    if document.mime_type == 'application/pdf':
        keyboard = [
            [InlineKeyboardButton("Extract Text", callback_data="Text")],
            [InlineKeyboardButton("Convert to DOCX", callback_data="DOCX")]
        ]
    elif document.mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        keyboard = [[InlineKeyboardButton("Convert to PDF", callback_data="PDF")]]
    else:
        await update.message.reply_text("Unsupported file type.")
        return

    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text("Choose an option:", reply_markup=reply_markup)

async def button_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle button callbacks."""
    query = update.callback_query
    await query.answer()

    # Check if text is available
    text = context.user_data.get("text")
    print(f"Text for conversion: {text}")  # Debugging line
    
    if not text:
        await query.message.reply_text("No text available for conversion. Please send a text first.")
        return

    try:
        if query.data == "Speech":
            # Convert text to speech
            output_path = os.path.join(tempfile.gettempdir(), "output.mp3")
            text_to_speech(text, output_path)
            await query.message.reply_audio(open(output_path, 'rb'))
            cleanup(output_path)
            
        elif query.data == "PDF":
            # Convert text to PDF
            output_path = os.path.join(tempfile.gettempdir(), "output.pdf")
            text_to_pdf(text, output_path)
            await query.message.reply_document(open(output_path, 'rb'))
            cleanup(output_path)
            
        elif query.data == "DOCX":
            # Convert text to DOCX
            output_path = os.path.join(tempfile.gettempdir(), "output.docx")
            doc = Document()
            doc.add_paragraph(text)
            doc.save(output_path)
            await query.message.reply_document(open(output_path, 'rb'))
            cleanup(output_path)

        else:
            await query.message.reply_text("Unsupported conversion.")
    
    except Exception as e:
        await query.message.reply_text(f"An error occurred: {e}")


# Add handlers
app.add_handler(CommandHandler("start", start))
app.add_handler(MessageHandler(filters.TEXT, handle_text))
app.add_handler(MessageHandler(filters.PHOTO, handle_image))
app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
app.add_handler(CallbackQueryHandler(button_callback))

if __name__ == "__main__":
    print("Bot is starting...")
    app.run_polling()
